"""TRPG document management: text extraction + chunking + vector retrieval.

Ported from ``nekro_trpg_dice_plugin``'s ``core/document_manager.py`` per the
M1 spec (``docs/specs/M1.md`` §4). ``DocumentProcessor`` is copied verbatim
(pure text parsing/chunking, no external services involved — only its type
hints were modernized to this repo's ``list[str]``/``X | None`` style).
``VectorDatabaseManager`` is re-pointed from the source's Qdrant client +
``gen_openai_embeddings`` calls onto this repo's embedded
``infra.vector.VectorStore`` + ``infra.embeddings.Embeddings`` abstractions;
every public method's name and argument shape, and the stored payload schema
(``{document_id, filename, chunk_index, text, chat_key, document_type,
created_at}``), are unchanged. Two behavioral differences from the source,
both required by the M1 spec:

- point ids are now deterministic (``f"{document_id}:{chunk_index}"``)
  instead of random UUIDs, since ``VectorStore.upsert`` is a natural
  upsert-by-id operation (Qdrant needed a fresh id per insert);
- ``answer_question`` now actually calls the injected ``LLMClient`` and
  returns its answer, instead of only building the RAG prompt and handing it
  back as a string for an outer framework to send (the source had no LLM
  client of its own to call; this port does, via the optional ``llm=``
  constructor argument).

PDF/DOCX parsing libraries (``pypdf`` / ``python-docx``) are optional
dependencies not installed by default in this repo: their imports are
guarded so the module always imports cleanly, TXT extraction always works,
and PDF/DOCX extraction raise a clear localized ``ValueError`` naming the
missing library instead of failing with an ``ImportError`` deep inside a
dependency.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from infra.embeddings import Embeddings
from infra.i18n import I18n, t
from infra.llm import LLMClient
from infra.vector import VectorStore

try:
    import pypdf

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx  # python-docx

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# Chunk break points, in priority order (paragraph break beats sentence-ending
# punctuation beats a bare newline beats a comma/semicolon). These are
# structural text-processing data, not user-visible UI text, so — like
# `core.prompt_sections`'s `_DOCUMENT_TYPE_EMOJI` — they stay literal instead
# of going through i18n.
_CHUNK_BREAK_POINTS = ("\n\n", "。", "！", "？", "\n", "，", "；")

# `delete_document` fetches every chunk of one document in a single `scroll`
# call; the embedded `VectorStore` has no pagination cursor (unlike Qdrant),
# so this just needs to be comfortably above any realistic per-document chunk
# count.
_DELETE_SCROLL_LIMIT = 100_000


def document_point_id(document_id: str, chunk_index: int) -> str:
    """Return the stable vector id shared by document storage and backup restore."""
    return f"{document_id}:{chunk_index}"


class DocumentProcessor:
    """Document parser: TXT/PDF/DOCX text extraction + character-based chunking.

    All methods are `staticmethod`s (pure text processing over in-memory
    bytes/str, no instance state).

    Supported formats:
    - TXT: plain text
    - MD/Markdown: plain text
    - PDF: requires the optional `pypdf` library
    - DOCX: requires the optional `python-docx` library
    """

    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from a TXT file, trying common encodings in turn."""
        try:
            for encoding in ("utf-8", "gbk", "gb2312", "big5"):
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # All strict decodes failed: fall back to error-tolerant utf-8.
            return file_content.decode("utf-8", errors="ignore")
        except Exception as e:
            raise ValueError(t("document.error.txt_parse_failed", error=str(e))) from e

    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from a PDF file (requires the optional `pypdf` library)."""
        if not PDF_AVAILABLE:
            raise ValueError(t("document.error.pdf_unavailable"))

        try:
            pdf_stream = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_stream)

            text_content = [page.extract_text() for page in pdf_reader.pages]
            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(t("document.error.pdf_parse_failed", error=str(e))) from e

    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from a DOCX file (requires the optional `python-docx` library)."""
        if not DOCX_AVAILABLE:
            raise ValueError(t("document.error.docx_unavailable"))

        try:
            docx_stream = io.BytesIO(file_content)
            doc = docx.Document(docx_stream)

            text_content = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # Table cell text is extracted too, one "|"-joined line per row.
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)
        except Exception as e:
            raise ValueError(t("document.error.docx_parse_failed", error=str(e))) from e

    @staticmethod
    def extract_text_by_extension(filename: str, file_content: bytes) -> str:
        """Dispatch to the right `extract_text_from_*` based on `filename`'s extension."""
        extension = filename.lower().split(".")[-1]

        if extension in ("txt", "md", "markdown"):
            return DocumentProcessor.extract_text_from_txt(file_content)
        elif extension == "pdf":
            return DocumentProcessor.extract_text_from_pdf(file_content)
        elif extension in ("docx", "doc"):
            return DocumentProcessor.extract_text_from_docx(file_content)
        else:
            raise ValueError(t("document.error.unsupported_format", extension=extension))

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 800) -> list[str]:
        """Split `text` into `chunk_size`-character pieces with `overlap`-character
        overlap between consecutive chunks.

        Each chunk boundary (other than the final one) is nudged back to the
        nearest `_CHUNK_BREAK_POINTS` match within the window, so chunks
        don't split mid-sentence when a clean break point is available.
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            if end < len(text):
                for break_char in _CHUNK_BREAK_POINTS:
                    break_pos = text.rfind(break_char, start, end)
                    if break_pos > start:
                        end = break_pos + len(break_char)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Next chunk's start, folding back by `overlap` (but never before `end`).
            start = max(start + chunk_size - overlap, end)

        return [chunk for chunk in chunks if chunk]


class VectorDatabaseManager:
    """Vector-backed document store: chunk + embed + upsert on `store_document`,
    cosine-similarity retrieval on `search_documents` / `get_document_context`.

    Built on an injected `infra.embeddings.Embeddings` + `infra.vector.VectorStore`
    pair instead of the source's Qdrant client + `gen_openai_embeddings` calls.
    `collection_name` is accepted for API/config-shape parity with the source
    but is not otherwise used: the embedded `VectorStore` has no notion of
    named collections, so callers construct one dedicated `VectorStore` (and
    therefore one dedicated `VectorDatabaseManager`) per logical collection.
    """

    def __init__(
        self,
        embeddings: Embeddings,
        vector_store: VectorStore,
        i18n: I18n,
        *,
        collection_name: str = "trpg_documents",
        max_search_results: int = 15,
        llm: LLMClient | None = None,
    ) -> None:
        self.embeddings = embeddings
        self.vector_store = vector_store
        self.i18n = i18n
        self.collection_name = collection_name
        self.max_search_results = max_search_results
        self.llm = llm
        self.document_processor = DocumentProcessor()

    async def store_document(
        self,
        document_id: str,
        filename: str,
        text_content: str,
        chat_key: str,
        document_type: str = "module",
    ) -> int:
        """Chunk, embed and upsert `text_content` into the vector store.

        Returns the number of chunks stored. Point ids are deterministic
        (``f"{document_id}:{chunk_index}"``), so re-storing the same
        `document_id` overwrites its previously stored chunks at matching
        indices rather than duplicating them.
        """
        chunks = self.document_processor.chunk_text(text_content)
        vectors = await self.embeddings.embed(chunks) if chunks else []
        created_at = datetime.now().isoformat()

        points = [
            (
                document_point_id(document_id, index),
                vector,
                {
                    "document_id": document_id,
                    "filename": filename,
                    "chunk_index": index,
                    "text": chunk,
                    "chat_key": chat_key,
                    "document_type": document_type,
                    "created_at": created_at,
                },
            )
            for index, (chunk, vector) in enumerate(zip(chunks, vectors, strict=True))
        ]

        # A backup produced by older versions may contain the same logical chunk
        # under a legacy namespaced id. Publish the canonical points first, then
        # remove every stale id for this room/document. This also removes trailing
        # chunks when a document is replaced with shorter content without risking
        # loss of the old document if the new upsert itself fails.
        existing = await self.vector_store.scroll(
            filter={"document_id": document_id, "chat_key": chat_key},
            limit=_DELETE_SCROLL_LIMIT,
        )
        canonical_ids = {point_id for point_id, _vector, _payload in points}
        stale_ids = [hit.id for hit in existing if hit.id not in canonical_ids]

        if points:
            await self.vector_store.upsert(points)
        if stale_ids:
            await self.vector_store.delete(stale_ids)

        return len(points)

    async def search_documents(
        self,
        query: str,
        chat_key: str,
        document_type: str | None = None,
        limit: int = 5,
    ) -> list[dict[str, Any]]:
        """Cosine-similarity search over `chat_key`'s documents, optionally
        narrowed to `document_type`.

        Returns dicts shaped like the source's Qdrant hits:
        `{id, score, text, filename, document_id, document_type, chunk_index}`.
        """
        [query_vector] = await self.embeddings.embed([query])

        search_filter: dict[str, Any] = {"chat_key": chat_key}
        if document_type:
            search_filter["document_type"] = document_type

        hits = await self.vector_store.search(query_vector, limit=limit, filter=search_filter)

        return [
            {
                "id": hit.id,
                "score": hit.score,
                "text": hit.payload["text"],
                "filename": hit.payload["filename"],
                "document_id": hit.payload["document_id"],
                "document_type": hit.payload["document_type"],
                "chunk_index": hit.payload["chunk_index"],
            }
            for hit in hits
        ]

    async def delete_document(self, document_id: str, chat_key: str) -> bool:
        """Delete every stored chunk of `document_id` scoped to `chat_key`.

        Returns True on success, including when there was nothing to delete
        (mirrors the source's behavior); returns False only if the store
        itself raises.
        """
        try:
            matches = await self.vector_store.scroll(
                filter={"document_id": document_id, "chat_key": chat_key},
                limit=_DELETE_SCROLL_LIMIT,
            )
            point_ids = [hit.id for hit in matches]
            if not point_ids:
                return True

            await self.vector_store.delete(point_ids)
            return True
        except Exception:
            return False

    async def list_documents(self, chat_key: str, document_type: str | None = None) -> list[dict[str, Any]]:
        """List each stored document once (its first chunk, `chunk_index == 0`,
        stands in for the whole document), scoped to `chat_key`."""
        list_filter: dict[str, Any] = {"chat_key": chat_key, "chunk_index": 0}
        if document_type:
            list_filter["document_type"] = document_type

        hits = await self.vector_store.scroll(filter=list_filter, limit=100)

        documents = []
        for hit in hits:
            text = hit.payload["text"]
            preview = text[:100] + "..." if len(text) > 100 else text
            documents.append(
                {
                    "document_id": hit.payload["document_id"],
                    "filename": hit.payload["filename"],
                    "document_type": hit.payload["document_type"],
                    "created_at": hit.payload["created_at"],
                    "preview": preview,
                }
            )
        return documents

    async def list_all_chunks(self, chat_key: str, limit: int = 1000, offset: Any = None) -> list[dict[str, Any]]:
        """List every stored chunk (full payload plus its `id`) for `chat_key`,
        used by module initialization to reassemble a document's full text.

        `offset` is accepted for source-API compatibility but unused: the
        embedded `VectorStore.scroll` has no native pagination cursor (it
        just returns up to `limit` matches in one call).
        """
        hits = await self.vector_store.scroll(filter={"chat_key": chat_key}, limit=limit)

        chunks = []
        for hit in hits:
            payload = dict(hit.payload)
            payload["id"] = hit.id
            chunks.append(payload)
        return chunks

    async def get_document_context(self, query: str, chat_key: str, max_context_length: int = 8000) -> str:
        """Concatenate the top matching chunks for `query` into one context blob,
        stopping before the next chunk would push the total past `max_context_length`
        characters."""
        search_results = await self.search_documents(query, chat_key, limit=self.max_search_results)

        if not search_results:
            return ""

        context_parts = []
        current_length = 0

        for result in search_results:
            chunk_info = self.i18n.t(
                "document.context.chunk_header", filename=result["filename"], text=result["text"]
            )

            if current_length + len(chunk_info) > max_context_length:
                break

            context_parts.append(chunk_info)
            current_length += len(chunk_info)

        return "\n".join(context_parts)

    async def answer_question(self, question: str, chat_key: str) -> str:
        """Answer `question` from `chat_key`'s uploaded documents via `self.llm`.

        Unlike the source (which only assembled the RAG prompt and left the
        actual model call to an outer Agent framework), this asks `self.llm`
        directly and returns its text answer, since an `LLMClient` is now
        injected alongside the vector store. Returns a localized "nothing
        found" message when no context matches, and a localized "no LLM
        configured" message if `self.llm` was never provided.
        """
        context = await self.get_document_context(question, chat_key, max_context_length=12000)

        if not context.strip():
            return self.i18n.t("document.answer.no_context")

        if self.llm is None:
            return self.i18n.t("document.answer.no_llm")

        prompt = self.i18n.t("document.answer.prompt", question=question, context=context)
        result = await self.llm.chat([{"role": "user", "content": prompt}])
        return result.content or ""
